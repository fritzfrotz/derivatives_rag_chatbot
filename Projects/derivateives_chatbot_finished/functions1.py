# Standard Library Imports
import os
from typing import List

# Third-Party Imports
from dotenv import load_dotenv
from docx import Document as DocxDocument

# LangChain Imports
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangChain Plugins
from langchain_chroma import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.output_parsers import StrOutputParser


# Load environment variables from .env file
load_dotenv()

# Access environment variables
tracing = os.getenv("LANGCHAIN_TRACING_V2")
endpoint = os.getenv("LANGCHAIN_ENDPOINT")
langchain_api_key = os.getenv("LANGCHAIN_API_KEY")
openai_api_key = os.getenv("OPENAI_API_KEY")

if not langchain_api_key or not openai_api_key:
    raise ValueError("API keys for LangChain and OpenAI are not set. Please check your .env file.")

# 1. Load Word documents using a relative path

def load_word_documents(relative_folder_path: str) -> List[Document]:
    """
    Loads Word documents from a specified folder and converts them into LangChain Document objects.
    Returns:
        List[Document]: Loaded DOCX documents represented as LangChain Document objects.
    """
    docs = []
    print(f"Loading documents from directory: {relative_folder_path}")

    for filename in os.listdir(relative_folder_path):
        if filename.endswith(".docx"):
            doc_path = os.path.join(relative_folder_path, filename)
            doc = DocxDocument(doc_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])

            if full_text.strip():  # Ensure there is non-whitespace content
                docs.append(Document(page_content=full_text, metadata={"source": filename}))
            else:
                print(f"Warning: {filename} is empty or contains only whitespace.")

    if not docs:
        print(f"Warning: No documents were loaded from {relative_folder_path}.")
    else:
        print(f"Loaded {len(docs)} documents from {relative_folder_path}")

    return docs


relative_folder_path = "./data/"
docs = load_word_documents(relative_folder_path)

# 2. Set up the text splitter

def split_text(documents: List[Document]) -> List[Document]:
    """
    Split the text content of the given list of Document objects into smaller chunks.

    Args:
        documents (List[Document]): List of Document objects containing text content to split.

    Returns:
        List[Document]: List of Document objects representing the split text chunks.
    """
    if not documents:
        print("No documents provided for splitting.")
        return []

    print(f"Splitting {len(documents)} documents into chunks.")

    # Initialize text splitter with specified parameters
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,  # Adjust this value as needed
        chunk_overlap=200,  # Overlap between consecutive chunks
        length_function=len,  # Function to compute the length of the text
        add_start_index=True,  # Flag to add start index to each chunk
    )

    # Split documents into smaller chunks using text splitter
    chunks = text_splitter.split_documents(documents)
    if not chunks:
        print("Warning: No chunks were generated after splitting the documents.")
    else:
        print(f"Split {len(documents)} documents into {len(chunks)} chunks.")

    return chunks  # Return the list of split text chunks


split_docs = split_text(docs)

# 3. Set up the vector store and retriever with persistent storage

persist_directory = "./chroma_store"  # Directory to save the Chroma vector store
absolute_path = os.path.abspath(persist_directory)  # Convert to absolute path
print(f"Chroma vector store will be saved in: {absolute_path}")

if not os.path.exists(persist_directory):
    os.makedirs(persist_directory)

# Check if a Chroma vector store already exists in the directory
if os.listdir(persist_directory):
    # Load existing Chroma vector store from disk
    vector_store = Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
else:
    # Create a new Chroma vector store from documents and save it to disk
    embeddings = OpenAIEmbeddings()
    vector_store = Chroma.from_documents(split_docs, embeddings, persist_directory=persist_directory)
    vector_store.persist()  # Save the vector store to disk

retriever = vector_store.as_retriever()


# 4. Define the template for the prompt that will generate the final answer

template = """You are a market finance assistant for question-answering tasks.
Below is the history of our conversation followed by some context retrieved from documents.

Use the following pieces of retrieved context to answer the question.
If the context does not provide a clear answer, refer to the conversation history to construct your response.
If you still don't know the answer, just say that you are a marked finance chatbot for derevatives academy. Keep the answer concise.

Note: The Conversation History is presented in chronological order, from the oldest to the most recent interactions.

Retrieved Context:
{context}

Conversation History:
{history}

Question: {question}
"""

prompt = ChatPromptTemplate.from_template(template)

# 5. Multi Query Template: Generate multiple perspectives for a given question
template_multiquery = """You are an AI language model assistant. Your task is to generate five
different versions of the given user question to retrieve relevant documents from a vector
database. By generating multiple perspectives on the user question, your goal is to help
the user overcome some of the limitations of the distance-based similarity search.
Provide these alternative questions separated by newlines. Original question: {question}"""
prompt_perspectives = ChatPromptTemplate.from_template(template_multiquery)

# Define the pipeline for generating multiple queries

generate_queries = (
    prompt_perspectives
    | ChatOpenAI(temperature=0)
    | StrOutputParser()
    | (lambda x: x.split("\n"))
)

# 6. Helper function to get unique documents
def get_unique_union(documents: list[list]):
    """Unique union of retrieved documents."""
    # Flatten list of lists, and convert each Document to string
    flattened_docs = [dumps(doc) for sublist in documents for doc in sublist]
    # Get unique documents
    unique_docs = list(set(flattened_docs))
    # Return as Document objects
    return [loads(doc) for doc in unique_docs]


# 7. Set up the LLM model
llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)
from langchain.memory import ConversationBufferWindowMemory

# Initialize memory to store only the last 5 interactions
memory = ConversationBufferWindowMemory(k=5)

def retrieve_and_generate_answer(question):
    """Retrieves relevant documents, includes recent history, and generates an answer."""

    # Step 1: Simple retrieval method (no multi-query)
    docs_simple = retriever.get_relevant_documents(question)

    if docs_simple:  # If documents were found
        context = docs_simple
    else:  # If no documents were found, use multi-query method
        generated_queries = generate_queries.invoke({"question": question})
        retrieved_docs_multi = [
            doc for query in generated_queries
            for doc in retriever.get_relevant_documents(query)
        ]
        context = get_unique_union(retrieved_docs_multi)

    # Combine the context into a single string for the prompt
    context_text = "\n\n".join([doc.page_content for doc in context])

    #  Load memory history and include the last interactions
    history = memory.load_memory_variables({})["history"]

    # Check if the history contains enough information
    if not history.strip():
        history = "This is the beginning of our conversation. There is no prior history."

    # Construct the prompt to include history, context, and the current question
    final_input = {
        "context": context_text,
        "question": question,
        "history": history
    }

    formatted_prompt = prompt.format(**final_input)

    # Generate the answer using the LLM
    result = llm(formatted_prompt)
    response_text = result.content  # Assuming result has a 'content' attribute

    # ** Save the new interaction to memory **
    memory.save_context({"input": question}, {"output": response_text})

    # Extract the list of document sources
    document_sources = [doc.metadata["source"] for doc in context]

    return response_text, document_sources


from datetime import datetime

def save_feedback(feedback, question, relative_folder_path, bot_response=None):
    """Save the user's feedback along with the time, initial question, and bot's response to a text file in the specified folder."""
    # Ensure the directory exists
    os.makedirs(relative_folder_path, exist_ok=True)

    # Define the feedback file path
    feedback_file = os.path.join(relative_folder_path, "feedback.txt")

    # Get the current time as a formatted string
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Format the feedback entry
    feedback_entry = f"Time: {current_time}\nQuestion: {question}\nBot Response: {bot_response}\nFeedback: {feedback}\n{'-'*40}\n"

    # Write the feedback entry to the file
    with open(feedback_file, "a") as f:
        f.write(feedback_entry)

def save_conversation(question, bot_response, relative_folder_path):
    """Save the question and the bot's response along with the time to a text file in the specified folder."""
    # Ensure the directory exists
    os.makedirs(relative_folder_path, exist_ok=True)

    # Define the conversation file path
    conversation_file = os.path.join(relative_folder_path, "conversation_log.txt")

    # Get the current time as a formatted string
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Format the conversation entry
    conversation_entry = f"Time: {current_time}\nQuestion: {question}\nBot Response: {bot_response}\n{'-'*40}\n"

    # Write the conversation entry to the file
    with open(conversation_file, "a") as f:
        f.write(conversation_entry)
